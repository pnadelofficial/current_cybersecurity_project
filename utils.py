from docx import Document
import re
import pandas as pd
import plotly.express as px
import os
import streamlit as st
import collections

# open expnder by default

codes = [
    "liberalism", 
    "constructivism", 
    "realism", 
    "cyberpersistence", 
    "policy_engineering_tasks"
]

color_mapping = {
    "Liberalism":"darkgreen", 
    "Constructivism":"darkblue", 
    "Realism":"darkred", 
    "Cyberpersistence":"saddlebrown", 
    "Policy engineering tasks":"teal"
}

color_mapping_lower = {
    "liberalism":"darkgreen", 
    "constructivism":"darkblue", 
    "realism":"darkred", 
    "cyberpersistence":"saddlebrown", 
    "policy_engineering_tasks":"teal"
}

double_codes_color_mapping = {
    "liberalism": "greens",
    "constructivism": "blues",
    "realism": "rdgy",
    "cyberpersistence":"brwnyl"
}

top_level = [
    'Deterrence.docx',
    'NLI.docx',
    'Persistence.docx',
    'International Norms.docx',
]

case_studies = {
    '2011':
        [
            '2010 National Security Strategy',
            '2011 DoD Cyber Strategy',
            '2011 National Military Strategy',
            '2011 International Strategy for Cyberspace',
            '2010 Quadrennial Defense Review'
        ],
    "2015":
        [
            "2015 National Security Strategy",
            '2015 White House Report on Cyber Deterrence Policy',
            '2014 Quadrennial Defense Review',
            '2015 DoD Cyber Strategy',
            '2015 National Military Strategy'
        ],
    '2018':
        [
            "2017 National Security Strategy",
            "2018 DoD Cyber Strategy Summary",
            "2018 National Cyber Strategy",
            "2018 National Defense Strategy Summary",
            "2018 National Military Strategy Description"
        ]
}

all_docs = []
for year, docs in case_studies.items():
    all_docs += docs

def read_doc_collect_codes(path):
    code = path.split('/')[-1].replace('.docx', '') 
    if path.endswith(".DS_Store"):
        return {}
    doc = Document(path)
    doc_text = '\n'.join([p.text for p in doc.paragraphs])
    doc_names = re.findall(r'Files\\\\20\d{2} Case Study\\\\.*Primary Sources_Policy_Strategies\\\\(.*) - .*', doc_text)
    doc_refs = re.split(r'Files\\\\20\d{2} Case Study\\\\.*Primary Sources_Policy_Strategies\\\\.*',doc_text)[1:]
    docs = list(zip(doc_names, doc_refs))

    doc_dict = {} 
    for name, dr in docs:
        refs = re.findall(r'Reference \d+ - .* Coverage\n(.*)', dr.strip())
        doc_dict[name] = [(code, r) for r in refs]

    return doc_dict

@st.cache_data
def get_pet_dict():
    pet_dict = collections.defaultdict(list)
    for file in os.listdir('./data/code_docs/policy_engineering_tasks'):
        if file == ".DS_Store":
            continue
        base_path = f'./data/code_docs/policy_engineering_tasks/{file}'
        for path in [f'{base_path}/{f}' for f in os.listdir(base_path) if 'DS_Store' not in f]:
            _dict = read_doc_collect_codes(path)
            for k, v in _dict.items():
                pet_dict[k] += v
    return pet_dict

def line_break(amt=2):
    for _ in range(amt):
        st.markdown("<br>", unsafe_allow_html=True)

def get_year_of_ps(ps):
    for year, docs in case_studies.items():
        if ps in docs:
            return year
    return None

class CaseStudy:
    def __init__(self, year): # primary_sources
        self.year = year
        if isinstance(year, list):
            self.primary_sources = []   
            for year in self.year:
                self.primary_sources += case_studies[year] # primary_sources
        else:
            self.primary_sources = case_studies[year]
    
    def _get_number_of_codes_by_primary_source(self, ps=None, code_choice=None):
        """For hierachy maps and bar charts"""
        results = []
        for code in codes:
            if code_choice and (code_choice != code):
                continue
            for root, dirs, files in os.walk(f'./data/code_docs/{code}/'):
                for file in files:
                    if file in ["Design.docx", "Maintenance.docx", "National Interests.docx", "Objective.docx", "Review.docx", "Strategy.docx"]:
                        continue
                    if not ps:
                        if file.endswith('.docx') and (file not in top_level):
                            cd = CodeDoc(os.path.join(root, file), self.year, purpose='hierarchy_map')
                            codes_by_ps = cd()
                            for _, tup in codes_by_ps.items():
                                results.append((cd.ancestor, cd.parent, cd.child, tup[-1]))
                    if file.endswith('.docx') and (file not in top_level): # deal with top level codes
                        cd = CodeDoc(os.path.join(root, file), self.year, purpose='hierarchy_map')
                        codes_by_ps = cd()
                        total = sum([tup[-1] for title, tup in codes_by_ps.items() if title == ps])
                        results.append((cd.ancestor, cd.parent, cd.child, total))
        return results
    
    def get_ca_dict(self, path):
        ca_dict = collections.defaultdict(list)
        for file in os.listdir(path):
            doc_path = f'{path}/{file}'
            cd = CodeDoc(doc_path, self.year, purpose='heat_map')
            _dict = cd()
            for k, v in _dict.items():
                ca_dict[k.strip()] += v
        return ca_dict
    
    def plot_heatmap(self, pet_dict, choice):
        ca_dict = self.get_ca_dict(f'./data/code_docs/{choice}/core_assumptions')
        docs = list(pet_dict.keys())
        color_scale = double_codes_color_mapping[choice]

        figs = []
        for doc in docs:
            _pet = pet_dict[doc]    
            _ca = ca_dict[doc]
            hm_data = []
            for t in _pet:
                if (t[0] == 'Objective') or (t[0] == 'National Interests'): 
                    for s in _ca:
                        if t[1] == s[1]:
                            hm_data.append((t[0], s[0]))         
            hmdf = pd.DataFrame(hm_data, columns=['policy engineering task', f'{choice} core assumptions'])
            pair_counts = hmdf.groupby([f'{choice} core assumptions', 'policy engineering task']).size().reset_index(name='Count')
            hm_data = pair_counts.pivot(index=f'{choice} core assumptions', columns='policy engineering task', values='Count').fillna(0)
            if len(hm_data) == 0:
                continue
            fig = px.imshow(hm_data, title=doc, text_auto=True, color_continuous_scale=color_scale, height=600)
            figs.append(fig)
        return figs

    def plot_treemap(self, results, ps=None):
        df = pd.DataFrame(results, columns=['ancestor', 'parent', 'child', 'refs'])
        df['ancestor'] = df['ancestor'].str.capitalize().apply(lambda x: x if x != "Policy_engineering_tasks" else 'Policy engineering tasks')
        df['parent'] = df['parent'].str.capitalize().apply(lambda x: x if x != 'Nli' else 'NLI')
        df['child'] = df['child'].str.capitalize()
        if ps:
            fig = px.treemap(df[~(df.child.isnull())] , path=[px.Constant(ps), 'ancestor', 'parent', 'child'], values='refs', color='ancestor', color_discrete_sequence=['lightgrey'], color_discrete_map=color_mapping)
        else:
            fig = px.treemap(df[~(df.child.isnull())] , path=[px.Constant("Home"), 'ancestor', 'parent', 'child'], values='refs', color='ancestor', color_discrete_sequence=['lightgrey'], color_discrete_map=color_mapping)
        fig.data[0].texttemplate = "<br><br><em style='font-size:75px'>%{label}</em><br>Code count (normalized by number of files): %{value}<br>Description: %{customdata[0]}"
        return fig

    def plot_bar_chart(self, results, title=None, showlegend=True):
        df = pd.DataFrame(results, columns=['Top level code', 'parent', 'codes', 'references']) 
        df = df[df.references > 0]
        color = [color_mapping_lower[c] for c in df['Top level code'].unique()]
        df['Top level code'] = df['Top level code'].str.capitalize().apply(lambda x: x if x != "Policy_engineering_tasks" else 'Policy engineering tasks')
        fig = px.bar(df, x='references', y='codes', title=title, orientation='h', color='Top level code', color_discrete_sequence=color)
        fig.update_layout(yaxis={"dtick":1},margin={"t":100,"b":100},height=900, showlegend=showlegend)    
        return fig

class CodeDoc:
    def __init__(self, path, year, purpose):
        self.path = path
        self.year = year if isinstance(year, list) else [year]
        self.purpose = purpose
        self.doc = Document(path)
        self.doc_text = '\n'.join([p.text for p in self.doc.paragraphs]) 

    def _collect_primary_source_chunks(self):
        poss_chunks = re.finditer(r"Files\\\\(20\d{2})", self.doc_text)
        poss_chunks = [m for m in poss_chunks]

        chunks = []
        for i, m in enumerate(poss_chunks):
            if m.group(1) in self.year:
                this_span = m.span()
                if i < len(poss_chunks) - 1:
                    next_span = poss_chunks[i+1].span()
                else:
                    next_span = (len(self.doc_text), len(self.doc_text))
                chunks.append(self.doc_text[this_span[0]:next_span[0]])
        return chunks
    
    def _get_number_of_codes_per_doc(self, chunk):
        # keeping title for now, but should thinking about how to use it
        title_check = re.search(r'.*Primary Sources_Policy_Strategies\\\\(.*)-', chunk)
        if title_check:
            title = title_check.group(1)
        else:
            return None
        m = re.search(r'§ (\d+) references coded', chunk)
        if m:
            return title.strip(), int(m.group(1))
        return title.strip(), 0
    
    def _collect_codes_hierarchy_bar(self):
        self.ancestor = self.path.split('/')[-3]  
        self.parent = self.path.split('/')[-2]
        self.child = self.path.split('/')[-1].replace('.docx', '')
        
        ps_with_codes = {}
        chunks = self._collect_primary_source_chunks()
        for chunk in chunks:
            tup = self._get_number_of_codes_per_doc(chunk)
            if tup is None:
                continue
            title, num_codes = tup
            ps_with_codes[title] = (self.ancestor, self.parent, self.child, num_codes)
        return ps_with_codes
    
    def _collect_codes_heat_maps(self):
        code = self.path.split('/')[-1].replace('.docx', '')
        chunks = self._collect_primary_source_chunks()

        doc_dict = {} 
        for chunk in chunks:
            name_check = re.search(r'.*Primary Sources_Policy_Strategies\\\\(.*)-', chunk)
            if name_check:
                name = name_check.group(1)
            else:
                continue
            refs = re.findall(r'Reference \d+ - .* Coverage\n(.*)', chunk)
            doc_dict[name] = [(code, r) for r in refs]
        return doc_dict
    
    def __call__(self):
        if (self.purpose == 'hierarchy_map') or (self.purpose == 'bar_chart'):
            return self._collect_codes_hierarchy_bar()
        elif self.purpose == 'heat_map':
            return self._collect_codes_heat_maps()